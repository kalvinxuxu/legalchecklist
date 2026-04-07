"""
修订模式 Word 文档生成器

使用 python-docx + XML 操作实现 Track Changes 功能
支持插入（insert）和删除（delete）修订
"""
import io
import re
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


class RevisionDocGenerator:
    """修订模式 Word 文档生成器"""

    def __init__(self):
        pass

    def create_revision_document(
        self,
        original_file_path: str,
        revisions: List[Dict[str, Any]]
    ) -> bytes:
        """
        创建带修订模式的 Word 文档

        Args:
            original_file_path: 原始 Word 文件路径
            revisions: 修订列表
                [{
                    "paragraph_index": 5,  # 段落索引
                    "type": "replace",      # replace / insert / delete
                    "original_text": "xxx", # 原文本（用于定位）
                    "new_text": "yyy",      # 新文本
                    "author": "AI 助手",
                    "comment": "风险修改建议",
                    "risk_description": "风险说明"
                }]

        Returns:
            修改后的 Word 文档字节流
        """
        doc = Document(original_file_path)

        # 按段落索引分组修订
        revisions_by_para: Dict[int, List[Dict[str, Any]]] = {}
        for rev in revisions:
            idx = rev.get("paragraph_index", 0)
            if idx not in revisions_by_para:
                revisions_by_para[idx] = []
            revisions_by_para[idx].append(rev)

        # 处理每个段落的修订
        for para_idx, para_revisions in revisions_by_para.items():
            if para_idx >= len(doc.paragraphs):
                continue

            para = doc.paragraphs[para_idx]
            self._apply_revisions_to_paragraph(para, para_revisions)

        # 保存到字节流
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _apply_revisions_to_paragraph(
        self,
        para,
        revisions: List[Dict[str, Any]]
    ) -> None:
        """
        将修订应用到单个段落

        使用 Word Track Changes (修订模式):
        - 删除的内容用 <w:del> 包裹
        - 新增的内容用 <w:ins> 包裹
        - 建议/评论作为批注添加
        """
        para_xml = para._p

        for rev in revisions:
            rev_type = rev.get("type", "replace")
            original_text = rev.get("original_text", "")
            new_text = rev.get("new_text", "")
            author = rev.get("author", "AI 助手")
            comment = rev.get("comment") or rev.get("risk_description", "")
            risk_level = rev.get("risk_level", "")

            if not original_text:
                continue

            if rev_type == "delete":
                # 删除：标记原文为删除，添加建议批注
                self._mark_text_as_deleted(para, original_text, author, comment)
            elif rev_type == "insert":
                # 插入：在原文后添加新文本
                self._add_inserted_text(para, original_text, new_text, author, comment)
            else:
                # replace：删除原文，插入新文本，并添加建议批注
                self._replace_text_with_revision(
                    para, original_text, new_text, author, comment, risk_level
                )

    def _find_text_in_paragraph(
        self,
        para,
        search_text: str
    ) -> List[Tuple[OxmlElement, str]]:
        """
        在段落中查找包含指定文本的所有 run 元素

        Returns:
            [(run_element, run_text), ...]
        """
        results = []
        for run in para.runs:
            if search_text in run.text:
                results.append((run._r, run.text))
        return results

    def _mark_text_as_deleted(
        self,
        para,
        text: str,
        author: str,
        comment: str
    ) -> None:
        """
        将指定文本标记为删除（Track Changes）
        """
        # 查找包含文本的 run
        runs = self._find_text_in_paragraph(para, text)
        if not runs:
            return

        # 获取段落的 XML 元素
        para_xml = para._p

        for run_elem, run_text in runs:
            # 创建 <w:del> 元素
            del_elem = OxmlElement('w:del')
            del_elem.set(qn('w:id'), str(self._generate_revision_id()))
            del_elem.set(qn('w:author'), author)
            del_elem.set(qn('w:date'), self._get_current_time())

            # 创建 <w:r> 元素
            r_elem = OxmlElement('w:r')

            # 创建 <w:delText> 元素
            del_text_elem = OxmlElement('w:delText')
            del_text_elem.set(qn('xml:space'), 'preserve')
            del_text_elem.text = text
            r_elem.append(del_text_elem)

            del_elem.append(r_elem)

            # 替换原 run 元素
            run_elem.addnext(del_elem)
            run_elem.getparent().remove(run_elem)

        # 添加批注
        if comment:
            self._add_comment_to_paragraph(para, author, comment, risk_level="high")

    def _add_inserted_text(
        self,
        para,
        before_text: str,
        new_text: str,
        author: str,
        comment: str
    ) -> None:
        """
        在指定文本后添加插入内容
        """
        runs = self._find_text_in_paragraph(para, before_text)
        if not runs:
            # 如果找不到原文，直接在段落末尾添加
            self._insert_text_at_run(para.runs[-1]._r if para.runs else para._p, new_text, author, comment)
            return

        # 在最后一个匹配的 run 后插入
        last_run_elem = runs[-1][0]
        self._insert_text_at_run(last_run_elem, new_text, author, comment)

    def _insert_text_at_run(
        self,
        after_elem,
        text: str,
        author: str,
        comment: str
    ) -> None:
        """
        在指定元素后插入带修订标记的文本
        """
        # 创建 <w:ins> 元素
        ins_elem = OxmlElement('w:ins')
        ins_elem.set(qn('w:id'), str(self._generate_revision_id()))
        ins_elem.set(qn('w:author'), author)
        ins_elem.set(qn('w:date'), self._get_current_time())

        # 创建 <w:r> 元素
        r_elem = OxmlElement('w:r')

        # 创建 <w:t> 元素
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = text
        r_elem.append(t_elem)

        ins_elem.append(r_elem)

        # 插入到段落中
        after_elem.addnext(ins_elem)

        # 添加批注
        if comment:
            self._add_comment_to_element(after_elem, author, comment)

    def _replace_text_with_revision(
        self,
        para,
        original_text: str,
        new_text: str,
        author: str,
        comment: str,
        risk_level: str = ""
    ) -> None:
        """
        用修订模式替换文本
        1. 将原文标记为删除
        2. 在删除内容后插入新文本
        3. 添加批注说明
        """
        runs = self._find_text_in_paragraph(para, original_text)
        if not runs:
            return

        # 获取段落 XML
        para_xml = para._p

        # 处理每个匹配的 run
        for i, (run_elem, run_text) in enumerate(runs):
            is_first = (i == 0)
            is_last = (i == len(runs) - 1)

            # 1. 创建 <w:del> 包裹原文
            del_elem = OxmlElement('w:del')
            del_elem.set(qn('w:id'), str(self._generate_revision_id()))
            del_elem.set(qn('w:author'), author)
            del_elem.set(qn('w:date'), self._get_current_time())

            r_del = OxmlElement('w:r')
            del_text = OxmlElement('w:delText')
            del_text.set(qn('xml:space'), 'preserve')
            del_text.text = original_text
            r_del.append(del_text)
            del_elem.append(r_del)

            # 2. 创建 <w:ins> 包裹新文本
            ins_elem = OxmlElement('w:ins')
            ins_elem.set(qn('w:id'), str(self._generate_revision_id() + 1))
            ins_elem.set(qn('w:author'), author)
            ins_elem.set(qn('w:date'), self._get_current_time())

            r_ins = OxmlElement('w:r')
            t_ins = OxmlElement('w:t')
            t_ins.set(qn('xml:space'), 'preserve')
            t_ins.text = new_text if is_last else ""
            r_ins.append(t_ins)
            ins_elem.append(r_ins)

            # 3. 替换原 run 元素
            run_elem.addnext(del_elem)
            run_elem.addnext(ins_elem)
            run_elem.getparent().remove(run_elem)

        # 4. 添加批注
        if comment:
            self._add_comment_to_paragraph(para, author, comment, risk_level)

    def _add_comment_to_paragraph(
        self,
        para,
        author: str,
        comment: str,
        risk_level: str = ""
    ) -> None:
        """
        为段落添加批注（使用文字颜色标记代替真正的批注，避免复杂实现）
        在段落末尾添加带颜色标记的建议文本
        """
        # 风险等级颜色
        color_map = {
            "high": RGBColor(0xC6, 0x28, 0x28),   # 红色
            "medium": RGBColor(0xE6, 0x51, 0x00), # 橙色
            "low": RGBColor(0x15, 0x65, 0xC0),     # 蓝色
        }
        color = color_map.get(risk_level, RGBColor(0x6A, 0x1B, 0x9A))  # 紫色

        # 在段落末尾添加建议文本
        run = para.add_run()
        run.font.color.rgb = color
        run.font.size = None  # 继承默认大小

        label = f" 【{author}建议】{comment}"
        run.text = label

    def _add_comment_to_element(
        self,
        elem,
        author: str,
        comment: str
    ) -> None:
        """为元素添加简单批注"""
        # 找到父段落
        para = elem.getparent()
        while para is not None and para.tag != qn('w:p'):
            para = para.getparent()

        if para is not None:
            self._add_comment_to_paragraph(para, author, comment)

    def _generate_revision_id(self) -> int:
        """生成唯一的修订 ID"""
        import random
        return random.randint(100000, 999999)

    def _get_current_time(self) -> str:
        """获取当前时间（ISO 格式）"""
        from datetime import datetime
        return datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

    def apply_suggestions_to_document(
        self,
        original_file_path: str,
        suggestions: List[Dict[str, Any]]
    ) -> bytes:
        """
        将建议应用到文档，生成带修订的文档

        这是对外的主要接口

        Args:
            original_file_path: 原始文件路径
            suggestions: 建议列表，每条建议包含：
                - paragraph_index: 段落索引
                - type: "replace" | "insert" | "delete"
                - original_text: 原文本
                - new_text: 新文本
                - risk_description: 风险描述
                - author: 建议作者

        Returns:
            修改后的 Word 字节流
        """
        return self.create_revision_document(original_file_path, suggestions)


# 全局实例
revision_doc_generator = RevisionDocGenerator()
