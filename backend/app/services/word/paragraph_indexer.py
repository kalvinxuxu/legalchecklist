"""
Word 段落索引服务

将 Word 文档解析为带索引的段落列表，支持精确定位
"""
from typing import Dict, Any, List, Optional
from docx import Document
from docx.oxml import OxmlElement
import io


class WordParagraphIndexer:
    """Word 段落索引器"""

    def __init__(self):
        self._cache: Dict[str, Dict[str, Any]] = {}

    async def parse_word_paragraphs(
        self,
        file_path: str,
        contract_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        解析 Word 文档，返回带索引的段落列表

        Args:
            file_path: Word 文件路径
            contract_id: 合同 ID（用于缓存）

        Returns:
            {
                "paragraphs": [
                    {"index": 0, "text": "...", "style": "Normal"},
                    {"index": 1, "text": "...", "style": "Heading 1"},
                    ...
                ],
                "tables": [...],
                "total_chars": 12345,
                "total_paragraphs": 100
            }
        """
        if contract_id and contract_id in self._cache:
            return self._cache[contract_id]

        doc = Document(file_path)

        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # 只保留有内容的段落
                paragraphs.append({
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else "Normal"
                })

        # 提取表格
        tables = []
        for t_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                "index": t_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data
            })

        # 统计
        total_chars = sum(len(p["text"]) for p in paragraphs)

        result = {
            "paragraphs": paragraphs,
            "tables": tables,
            "total_chars": total_chars,
            "total_paragraphs": len(paragraphs)
        }

        if contract_id:
            self._cache[contract_id] = result

        return result

    def get_paragraph_by_index(
        self,
        file_path: str,
        index: int
    ) -> Optional[Dict[str, Any]]:
        """
        根据索引获取指定段落内容

        Args:
            file_path: Word 文件路径
            index: 段落索引

        Returns:
            段落信息，包含 text, style, runs
        """
        doc = Document(file_path)

        if index < 0 or index >= len(doc.paragraphs):
            return None

        para = doc.paragraphs[index]
        runs = []
        for run in para.runs:
            runs.append({
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline
            })

        return {
            "index": index,
            "text": para.text,
            "style": para.style.name if para.style else "Normal",
            "runs": runs
        }

    async def search_paragraphs(
        self,
        file_path: str,
        keyword: str,
        case_sensitive: bool = False
    ) -> List[Dict[str, Any]]:
        """
        搜索包含关键词的段落

        Args:
            file_path: Word 文件路径
            keyword: 搜索关键词
            case_sensitive: 是否区分大小写

        Returns:
            匹配的段落列表
        """
        parse_result = await self.parse_word_paragraphs(file_path)
        matches = []

        search_text = keyword if case_sensitive else keyword.lower()

        for para in parse_result["paragraphs"]:
            text = para["text"] if case_sensitive else para["text"].lower()
            if search_text in text:
                matches.append(para)

        return matches

    def invalidate_cache(self, contract_id: str) -> None:
        """清除指定合同的缓存"""
        if contract_id in self._cache:
            del self._cache[contract_id]


# 全局实例
word_indexer = WordParagraphIndexer()
