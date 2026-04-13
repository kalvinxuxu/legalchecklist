"""
阿里云文档智能解析服务
支持 PDF 和 Word 文件解析
"""
import httpx
import hmac
import hashlib
import base64
import urllib.parse
from typing import Optional, Dict, Any
from app.core.config import settings


class AliyunDocumentParser:
    """阿里云文档智能解析客户端"""

    def __init__(self):
        self.access_key_id = settings.ALIYUN_ACCESS_KEY_ID
        self.access_key_secret = settings.ALIYUN_ACCESS_KEY_SECRET
        self.endpoint = settings.ALIYUN_OSS_ENDPOINT
        self.bucket = settings.ALIYUN_OSS_BUCKET
        self.timeout = 60.0

        # 文档智能 API 端点
        self.docint_endpoint = "https://docint.cn-hangzhou.aliyuncs.com"

    def _generate_signature(self, method: str, path: str, params: dict) -> str:
        """生成阿里云 API 签名"""
        # 规范化的查询字符串
        sorted_params = sorted(params.items())
        canonicalized_query_string = "&".join(
            f"{k}={urllib.parse.quote(str(v), safe='')}" for k, v in sorted_params
        )

        # 签名字符串
        string_to_sign = f"{method}\n\n\n\n\n{canonicalized_query_string}"

        # HMAC-SHA1 签名
        h = hmac.new(
            self.access_key_secret.encode('utf-8'),
            string_to_sign.encode('utf-8'),
            hashlib.sha1
        )
        signature = base64.b64encode(h.digest()).decode('utf-8')

        return signature

    async def parse_pdf(
        self,
        file_path: str,
        ocr: bool = False
    ) -> Dict[str, Any]:
        """
        解析 PDF 文件

        Args:
            file_path: OSS 上的文件路径 或 本地文件路径
            ocr: 是否启用 OCR（扫描版 PDF 需要）

        Returns:
            解析结果，包含：
            - text: 提取的文本
            - pages: 页数
            - elements: 结构化元素（段落、标题、表格等）
        """
        # 方案 1：使用阿里云文档智能 API
        if self.access_key_id and self.access_key_secret:
            return await self._parse_with_aliyun(file_path, "pdf", ocr)

        # 方案 2：降级使用本地解析（pdfplumber）
        return await self._parse_pdf_locally(file_path)

    async def parse_word(
        self,
        file_path: str
    ) -> Dict[str, Any]:
        """
        解析 Word 文件

        Args:
            file_path: OSS 上的文件路径 或 本地文件路径

        Returns:
            解析结果，包含：
            - text: 提取的文本
            - paragraphs: 段落列表
            - tables: 表格列表
        """
        # 方案 1：使用阿里云文档智能 API
        if self.access_key_id and self.access_key_secret:
            return await self._parse_with_aliyun(file_path, "word", False)

        # 方案 2：降级使用本地解析（python-docx）
        return await self._parse_word_locally(file_path)

    async def _parse_with_aliyun(
        self,
        file_path: str,
        file_type: str,
        ocr: bool
    ) -> Dict[str, Any]:
        """使用阿里云文档智能 API 解析"""
        # TODO: 实现完整的阿里云文档智能 API 调用
        # 这里提供简化版本，实际使用时需要完善签名和请求逻辑

        async with httpx.AsyncClient(timeout=self.timeout) as client:
            # 构造请求参数
            params = {
                "Action": "ExtractText",
                "Version": "2021-07-07",
                "Timestamp": self._get_timestamp(),
                "Format": "JSON",
                "SignatureMethod": "HMAC-SHA1",
                "SignatureVersion": "1.0",
                "SignatureNonce": self._get_nonce(),
                "AccessKeyId": self.access_key_id,
            }

            # 添加文件类型特定参数
            if file_type == "pdf":
                params["FileType"] = "PDF"
                params["EnableOCR"] = "true" if ocr else "false"
            else:
                params["FileType"] = "DOCX"

            # 生成签名
            signature = self._generate_signature("GET", "/", params)
            params["Signature"] = signature

            # 调用 API
            try:
                response = await client.get(
                    f"https://docint.cn-hangzhou.aliyuncs.com",
                    params=params
                )
                response.raise_for_status()
                result = response.json()

                return {
                    "text": result.get("Text", ""),
                    "pages": result.get("PageCount", 1),
                    "elements": result.get("Elements", []),
                    "source": "aliyun"
                }
            except Exception as e:
                # API 调用失败，降级到本地解析
                print(f"阿里云文档智能 API 调用失败：{e}")
                if file_type == "pdf":
                    return await self._parse_pdf_locally(file_path)
                else:
                    return await self._parse_word_locally(file_path)

    async def _parse_pdf_locally(self, file_path: str) -> Dict[str, Any]:
        """
        本地解析 PDF（降级方案）

        优先级：pdfplumber > pymupdf > EasyOCR > Tesseract OCR
        """
        import os
        from pathlib import Path

        # 确保路径是绝对路径
        if not os.path.isabs(file_path):
            # 获取 backend 目录（uploads 在 backend 下，不在 app 下）
            backend_dir = Path(__file__).resolve().parent.parent.parent.parent
            file_path = str(backend_dir / file_path)

        text_parts = []
        pages = 0

        # 方法 1：使用 pdfplumber
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                pages = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)

            if text_parts and len("\n".join(text_parts).strip()) > 100:
                full_text = "\n\n".join(text_parts)
                # 检测是否为乱码（包含过多不可见字符或特殊字符）
                if self._is_garbled_text(full_text):
                    print(f"pdfplumber 提取的文本疑似乱码，尝试 OCR...")
                else:
                    return {
                        "text": full_text,
                        "pages": pages,
                        "elements": [],
                        "source": "local_pdfplumber"
                    }
        except Exception as e:
            print(f"pdfplumber 解析失败: {e}")

        # 方法 2：使用 pymupdf (fitz)
        text_parts = []
        try:
            import fitz

            doc = fitz.open(file_path)
            pages = len(doc)

            for page_num in range(pages):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)

            doc.close()

            if text_parts and len("\n".join(text_parts).strip()) > 100:
                full_text = "\n\n".join(text_parts)
                # 检测是否为乱码
                if self._is_garbled_text(full_text):
                    print(f"pymupdf 提取的文本疑似乱码，尝试 OCR...")
                else:
                    return {
                        "text": full_text,
                        "pages": pages,
                        "elements": [],
                        "source": "local_pymupdf"
                    }
        except Exception as e:
            print(f"pymupdf 解析失败: {e}")

        # 方法 3：使用 EasyOCR（扫描版 PDF）
        print("尝试使用 EasyOCR 进行 OCR 识别...")
        ocr_result = await self._parse_with_easyocr(file_path)
        if ocr_result.get("text"):
            return ocr_result

        # 方法 4：使用 Tesseract OCR（备用）
        print("尝试使用 Tesseract OCR 进行识别...")
        ocr_result = await self._parse_with_tesseract(file_path)
        if ocr_result.get("text"):
            return ocr_result

        return {
            "text": "",
            "pages": pages,
            "elements": [],
            "source": "local",
            "error": "无法从 PDF 提取文本，所有 OCR 方法均失败"
        }

    async def _parse_with_easyocr(self, file_path: str) -> Dict[str, Any]:
        """使用 EasyOCR 解析扫描版 PDF"""
        try:
            import fitz
            from PIL import Image
            import io
            import numpy as np
            import os
            import gc

            # 初始化 EasyOCR（支持中文+英文）
            import easyocr
            reader = easyocr.Reader(['ch_sim', 'en'], gpu=False, verbose=False)

            doc = fitz.open(file_path)
            pages = len(doc)
            all_text = []

            for page_num in range(pages):
                page = doc[page_num]
                # 将页面转换为图像（使用 1.5x 缩放减少内存）
                mat = fitz.Matrix(1.5, 1.5)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                img_array = np.array(img)

                # OCR 识别
                result = reader.readtext(img_array)
                if result:
                    page_texts = []
                    for detection in result:
                        # result 格式: (bbox, text, confidence)
                        if len(detection) >= 2:
                            text = detection[1]
                            if text and text.strip():
                                page_texts.append(text.strip())
                    if page_texts:
                        all_text.append("\n".join(page_texts))

                # 释放内存
                del img, img_array, result
                gc.collect()

                # 进度提示
                if (page_num + 1) % 5 == 0:
                    print(f"  EasyOCR 处理进度: {page_num + 1}/{pages}")

            doc.close()

            if all_text and len("\n".join(all_text).strip()) > 50:
                return {
                    "text": "\n\n".join(all_text),
                    "pages": pages,
                    "elements": [],
                    "source": "easyocr"
                }

        except MemoryError as e:
            print(f"EasyOCR 内存不足: {e}")
            return {"text": "", "pages": 0, "source": "easyocr", "error": "内存不足"}
        except Exception as e:
            print(f"EasyOCR 解析失败: {e}")

        return {"text": "", "pages": 0, "source": "easyocr"}

    async def _parse_with_tesseract(self, file_path: str) -> Dict[str, Any]:
        """使用 Tesseract OCR 解析扫描版 PDF"""
        try:
            import fitz
            from PIL import Image
            import pytesseract
            import io
            import os

            # 配置 tesseract 路径（Windows 安装路径）
            tesseract_exe = None
            tesseract_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                r"C:\Tesseract-OCR\tesseract.exe"
            ]
            for path in tesseract_paths:
                if os.path.exists(path):
                    tesseract_exe = path
                    pytesseract.pytesseract.tesseract_cmd = path
                    break

            if not tesseract_exe:
                print("Tesseract 未安装")
                return {"text": "", "pages": 0, "source": "tesseract"}

            # 检查语言包
            tessdata_path = os.path.join(os.path.dirname(tesseract_exe), 'tessdata')
            chi_sim_path = os.path.join(tessdata_path, 'chi_sim.traineddata')
            if not os.path.exists(chi_sim_path):
                print(f"Tesseract 中文语言包未找到: {chi_sim_path}")
                return {"text": "", "pages": 0, "source": "tesseract"}

            doc = fitz.open(file_path)
            pages = len(doc)
            all_text = []

            for page_num in range(pages):
                page = doc[page_num]
                # 将页面转换为图像
                mat = fitz.Matrix(1.5, 1.5)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))

                # OCR 识别（中文+英文）
                text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                if text and text.strip():
                    all_text.append(text.strip())

                if (page_num + 1) % 5 == 0:
                    print(f"  Tesseract OCR 处理进度: {page_num + 1}/{pages}")

            doc.close()

            if all_text and len("\n".join(all_text).strip()) > 50:
                return {
                    "text": "\n\n".join(all_text),
                    "pages": pages,
                    "elements": [],
                    "source": "tesseract"
                }

        except Exception as e:
            print(f"Tesseract OCR 解析失败: {e}")

        return {"text": "", "pages": 0, "source": "tesseract"}

    async def _parse_word_locally(self, file_path: str) -> Dict[str, Any]:
        """
        本地解析 Word（降级方案）

        使用 python-docx 库
        """
        try:
            from docx import Document
        except ImportError:
            return {
                "text": "",
                "error": "python-docx 未安装，请运行 pip install python-docx",
                "source": "local"
            }

        doc = Document(file_path)

        # 提取段落
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 提取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return {
            "text": "\n\n".join(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables,
            "source": "local_python_docx"
        }

    def _is_garbled_text(self, text: str) -> bool:
        """
        检测文本是否为乱码

        判断逻辑：
        1. 替换字符（U+FFFD �）比例过高
        2. 不可见字符（如 \x00-\x08）比例过高
        3. 连续 ASCII 字符（如 I、█）比例过高
        4. 中文字符比例为 0 且文本较长
        """
        if not text:
            return True

        # 统计各类字符
        total = len(text)
        replacement_char = '\ufffd'  # U+FFFD 替换字符
        replacement_count = text.count(replacement_char)
        invisible = sum(1 for c in text if ord(c) < 0x20 and c not in '\n\r\t')
        ascii_count = sum(1 for c in text if 32 <= ord(c) < 127 and c not in '\n\r\t ')
        chinese_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')

        # 替换字符超过 10%（PDF 字体无法识别的标志）
        if replacement_count / total > 0.1:
            return True

        # 不可见字符超过 5%
        if invisible / total > 0.05:
            return True

        # 纯 ASCII 字符（可能是乱码如 █████ 或 IIIIIII）
        if ascii_count / total > 0.7 and chinese_count == 0:
            return True

        # 中文字符为 0 且文本长度 > 50（正常中文文本应有中文）
        if chinese_count == 0 and total > 50:
            return True

        return False

    def _get_timestamp(self) -> str:
        """获取 ISO8601 格式时间戳"""
        from datetime import datetime, timezone
        return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    def _get_nonce(self) -> str:
        """生成随机数"""
        import uuid
        return str(uuid.uuid4())


# 全局解析器实例
document_parser = AliyunDocumentParser()
